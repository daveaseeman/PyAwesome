#!/home/vagrant/miniconda3/bin/python

import docx
from combine_data import combine_data


combined_data, temperature_fail, unknown_fail  = combine_data()

total_tests = len(combined_data)

tests_failed = len(temperature_fail) + len(unknown_fail)

test_passed = total_tests - tests_failed

test_fail_temp = len(temperature_fail) 

test_fail_unknown = len(unknown_fail)


document = docx.Document()

document.add_heading("Acme Test report", 0)

document.add_paragraph('')


test_ran_dates_str = "Tests Summary: Tests ran from: " + combined_data[0][0] + " to " + combined_data[-1][0]
document.add_heading(test_ran_dates_str, level  = 1)
document.add_paragraph('')

p = document.add_paragraph("Total number of tests: ")

p.add_run(str(total_tests)).bold = True


p = document.add_paragraph("Tests Passed: ")

p.add_run(str(test_passed)).bold = True

p = document.add_paragraph("Tests Failed: ")

p.add_run(str(tests_failed)).bold = True

p = document.add_paragraph("Tests Failed due to temperature: ")

p.add_run(str(test_fail_temp)).bold = True

p = document.add_paragraph("Tests Failed due to unknown causes: ")

p.add_run(str(test_fail_unknown)).bold = True

document.add_paragraph('')
document.add_heading("Table of Results ", level  = 2)
document.add_paragraph('')

table = document.add_table(rows = 1, cols = 3)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Test Date / Time'
hdr_cells[1].text = 'Result'
hdr_cells[2].text = 'Failure Cause'

for data in combined_data:
    row_cells = table.add_row().cells
    row_cells[0].text = data[0]
    row_cells[1].text = data[1]

    if data[0] in temperature_fail:
        row_cells[2].text = "Temperature Failure"
    elif data[0] in unknown_fail:
        row_cells[2].text = "Unknown Failure- Please investigate"
    else:
        row_cells[2].text = ""

document.save("Test Report.docx")



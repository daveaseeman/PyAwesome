#!/home/vagrant/miniconda3/bin/python

import docx

document = docx.Document("Chess Club Rules.docx")

print("Number of paragraphs = {}".format(len(document.paragraphs)))

print(document.paragraphs[0].text)

for para in document.paragraphs:
    print(para.text)
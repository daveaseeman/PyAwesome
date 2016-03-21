#!/usr/bin/python3

from docx import Document
import pdb

def extract_data(filename):

    data = []
    with open(filename, "r") as f:
        data = f.read()

    data2 = data.split("\n")

    data_return = []
    for i in range(2,len(data2)):
       data_return.append(data2[i].split(","))

    return data_return


# cpu         = extract_data("cpu.txt")
temperature        = extract_data("temperature.txt")
pass_result = extract_data("pass-fail.txt")

#print(cpu)
#print(temperature)
#print(pass_result)

combined_data = []

for i in range(len(temperature)):
    combined_data.append([temperature[i][0], temperature[i][1], pass_result[i][1]])

#print(combined_data[:2])

# pdb.set_trace()

#cpu_fail = []
temperature_fail = []
unknown_fail = []
total_fail = 0

for data in combined_data:
    if data[2] == "FAIL":
        total_fail += 1
        if int(data[1]) >= 35:
            temperature_fail.append(data[0])
        else:
            unknown_fail.append(data[0])

print("temperature fails: ", temperature_fail)
print("unknown_fail: ", unknown_fail)

document = Document()

document.add_heading('Test Report for Acme', 0)

doc_test_str = "Test Results for " + combined_data[0][0] + " to " + combined_data[-1][0]

document.add_heading(doc_test_str, level=1)

document.save('Test Results.docx')
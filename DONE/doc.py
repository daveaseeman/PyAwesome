#!/home/vagrant/miniconda3/bin/python

import docx

doc = docx.Document("Chess Club Rules.docx")

print("Number of paragraphs ", len(doc.paragraphs))

for i in range(8):
    print(doc.paragraphs[i].text)